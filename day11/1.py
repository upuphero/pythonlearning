#coding=utf-8
from docx import Document
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #chinese

#路径就是你文档所在的路径,word_name文档文件名称
document=Document(路径+'/'+word_name)

#old_text原先模板上的文字,new_text是将要替换的文字,document是需要处理的文档对象
def change_text(old_text, new_text,document):
    #遍历文档内所有段落
    all_paragraphs = document.paragraphs

    for paragraph in all_paragraphs:

        for run in paragraph.runs:

            if old_text in run.text:

                run.text = run.text.replace(old_text, new_text)
    #遍历文档内所有表格
    all_tables = document.tables

    for table in all_tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    for run in paragraph.runs:

                        #print '扫描-->'+run.text

                        if old_text in run.text:

                            #print run.text+'->'+new_text #+'-->'+chardet.detect(cell.text.encode())

                            run.text = run.text.replace(old_text, new_text)
