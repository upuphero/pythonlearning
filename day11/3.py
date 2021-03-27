from docx import Document
from docx.shared import Cm
# from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

document=Document(r"D:\file\download\1.docx")
# 首先对段落格式进行修改，docx默认标题也属于段落，因此“表扬信”是第一段
#paragraphs=document.paragraphs
#paragraphs[2].paragraph_format.first_line_indent=Cm(0.74)
#paragraphs[3].paragraph_format.left_indent=Cm(0.74)
#paragraphs[4].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
#paragraphs[4].paragraph_format.right_indent=Cm(2)
#paragraphs[5].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
#paragraphs[5].paragraph_format.right_indent=Cm(2)
# 对文本进行修改
# 修改第二段
#paragraphs[1].text="小Z同学："
# 将第三段陆亦可替换为大Z，她替换为他。通过python的正则表达式，可以很简单地实现文本的替换和查找。
for i in document.paragraphs:
    text=re.sub('1','他',text)
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                if key in para.runs[i].text:
                    print(key+"-->"+value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)
    return document

for i in
text=re.sub('陆亦可','大Z',paragraphs[2].text)
text=re.sub('她','他',text)
paragraphs[2].text=text
# 在第四段后面加上
paragraphs[3].add_run("向小z同学学习！")
# 修改表格里面的内容
#tables=document.tables
#tables[0].cell(1,0).text="猫粮"
#tables[0].cell(2,0).text="猫粮"
#tables[0].cell(3,0).text="猫粮"
# 插入一张图片，图片宽度设置为11.8cm
# document.add_picture('fun.jpg', width=Cm(11.8))
document.save()